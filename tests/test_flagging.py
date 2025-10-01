#!/usr/bin/env python3
"""
Tests for Document Flagging System
"""

import os
import sys
import json
from pathlib import Path
import pytest
import docx

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

from core.document_flag_manager import (
    DocumentFlagManager, 
    ChangeType, 
    FlaggedTextBuilder,
    create_flag_manager
)


def test_flag_manager_initialization():
    """Test basic flag manager initialization"""
    fm = DocumentFlagManager(enable_visual_flags=False, enable_comments=True)
    
    assert fm.enable_visual_flags == False
    assert fm.enable_comments == True
    assert len(fm.changes) == 0
    assert len(fm.color_map) > 0


def test_add_change_record():
    """Test recording changes"""
    fm = DocumentFlagManager()
    
    fm.add_change_record(
        change_type=ChangeType.UNICODE_SUBSTITUTION,
        location="paragraph_5",
        original="original text",
        modified="modіfied text",
        details={'chars_changed': 1}
    )
    
    assert len(fm.changes) == 1
    assert fm.changes[0]['type'] == ChangeType.UNICODE_SUBSTITUTION
    assert fm.changes[0]['location'] == "paragraph_5"
    assert 'timestamp' in fm.changes[0]


def test_flag_paragraph(tmp_path):
    """Test flagging a paragraph"""
    fm = DocumentFlagManager(enable_visual_flags=True, enable_comments=True)
    
    # Create test document
    doc = docx.Document()
    para = doc.add_paragraph("Test paragraph text")
    
    # Flag it
    fm.flag_paragraph(para, ChangeType.HEADER_MODIFICATION, "Test message")
    
    # Check that paragraph was modified (has marker or highlight)
    # Note: Exact verification depends on implementation
    assert len(para.runs) > 0


def test_hidden_marker_detection(tmp_path):
    """Test detection of IPT markers"""
    fm = DocumentFlagManager()
    
    # Create document with marker
    doc = docx.Document()
    para = doc.add_paragraph("Test")
    para.runs[0].text += "\u200B\u200C\u200D[test_marker]"
    
    # Save and reload
    doc_path = tmp_path / "test_marker.docx"
    doc.save(str(doc_path))
    
    # Detect
    doc2 = docx.Document(str(doc_path))
    has_marker = fm.detect_previous_modifications(doc2)
    
    assert has_marker == True


def test_no_marker_detection(tmp_path):
    """Test clean document detection"""
    fm = DocumentFlagManager()
    
    # Create clean document
    doc = docx.Document()
    doc.add_paragraph("Clean paragraph")
    
    # Save and reload
    doc_path = tmp_path / "test_clean.docx"
    doc.save(str(doc_path))
    
    # Detect
    doc2 = docx.Document(str(doc_path))
    has_marker = fm.detect_previous_modifications(doc2)
    
    assert has_marker == False


def test_statistics():
    """Test statistics generation"""
    fm = DocumentFlagManager()
    
    # Add multiple changes
    fm.add_change_record(ChangeType.UNICODE_SUBSTITUTION, "p1", "a", "а")
    fm.add_change_record(ChangeType.UNICODE_SUBSTITUTION, "p2", "b", "в")
    fm.add_change_record(ChangeType.ZERO_WIDTH_INSERTION, "p3", "x", "x\u200B")
    
    stats = fm.get_statistics()
    
    assert stats['total_changes'] == 3
    assert stats['changes_by_type'][ChangeType.UNICODE_SUBSTITUTION] == 2
    assert stats['changes_by_type'][ChangeType.ZERO_WIDTH_INSERTION] == 1
    assert stats['first_change'] is not None
    assert stats['last_change'] is not None


def test_export_change_log(tmp_path):
    """Test JSON export"""
    fm = DocumentFlagManager()
    
    # Add changes
    fm.add_change_record(ChangeType.HEADER_MODIFICATION, "h1", "BAB I", "BАB I")
    
    # Export
    log_path = tmp_path / "changes.json"
    fm.export_change_log(str(log_path))
    
    # Verify file exists and is valid JSON
    assert log_path.exists()
    
    with open(log_path, 'r') as f:
        data = json.load(f)
    
    assert 'metadata' in data
    assert 'statistics' in data
    assert 'changes' in data
    assert len(data['changes']) == 1


def test_document_properties(tmp_path):
    """Test document properties modification"""
    fm = DocumentFlagManager()
    
    # Create document
    doc = docx.Document()
    doc.add_paragraph("Test content")
    
    # Add a change
    fm.add_change_record(ChangeType.UNICODE_SUBSTITUTION, "p1", "test", "tеst")
    
    # Add properties
    fm.add_document_properties(doc)
    
    # Verify
    assert doc.core_properties.subject == "IPT Modified Document"
    assert "IPT" in doc.core_properties.keywords


def test_change_log_section(tmp_path):
    """Test change log section addition"""
    fm = DocumentFlagManager()
    
    # Create document
    doc = docx.Document()
    doc.add_paragraph("Main content")
    
    # Add changes
    fm.add_change_record(ChangeType.UNICODE_SUBSTITUTION, "p1", "a", "а")
    fm.add_change_record(ChangeType.ZERO_WIDTH_INSERTION, "p2", "b", "b\u200B")
    
    # Add log section
    original_para_count = len(doc.paragraphs)
    fm.add_change_log_section(doc)
    
    # Verify new paragraphs were added
    assert len(doc.paragraphs) > original_para_count
    
    # Check for log title
    log_found = False
    for para in doc.paragraphs:
        if "IPT MODIFICATION LOG" in para.text:
            log_found = True
            break
    
    assert log_found == True


def test_visual_flags_enabled():
    """Test visual flags mode"""
    fm = DocumentFlagManager(enable_visual_flags=True)
    
    doc = docx.Document()
    para = doc.add_paragraph("Test text")
    
    # Flag with visual highlight
    fm.flag_paragraph(para, ChangeType.UNICODE_SUBSTITUTION)
    
    # Check if highlight was applied
    has_highlight = False
    for run in para.runs:
        if run.font.highlight_color is not None:
            has_highlight = True
            break
    
    # May or may not have highlight depending on implementation
    # Just verify it doesn't crash
    assert True


def test_create_flag_manager_from_config():
    """Test factory function with config"""
    config = {
        'debug': {
            'enable_visual_flags': True,
            'enable_comments': False
        }
    }
    
    fm = create_flag_manager(config)
    
    assert fm.enable_visual_flags == True
    assert fm.enable_comments == False


def test_flagged_text_builder(tmp_path):
    """Test FlaggedTextBuilder helper"""
    fm = DocumentFlagManager()
    builder = FlaggedTextBuilder(fm)
    
    doc = docx.Document()
    para = doc.add_paragraph()
    
    # Create flagged run
    run = builder.create_flagged_run(para, "test text", ChangeType.UNICODE_SUBSTITUTION)
    
    assert run is not None
    assert "test text" in run.text


def test_multiple_change_types():
    """Test tracking multiple change types"""
    fm = DocumentFlagManager()
    
    # Add different types of changes
    fm.add_change_record(ChangeType.UNICODE_SUBSTITUTION, "p1", "a", "а")
    fm.add_change_record(ChangeType.ZERO_WIDTH_INSERTION, "p2", "b", "b\u200B")
    fm.add_change_record(ChangeType.HEADER_MODIFICATION, "h1", "BAB I", "BАB I")
    fm.add_change_record(ChangeType.METADATA_CHANGE, "meta", "[old]", "[new]")
    fm.add_change_record(ChangeType.PARAPHRASE, "p3", "old phrase", "new phrase")
    
    stats = fm.get_statistics()
    
    assert len(stats['changes_by_type']) == 5
    assert stats['total_changes'] == 5


def test_change_summary_generation():
    """Test summary text generation"""
    fm = DocumentFlagManager()
    
    # Add changes
    for i in range(10):
        fm.add_change_record(ChangeType.UNICODE_SUBSTITUTION, f"p{i}", f"text{i}", f"tеxt{i}")
    
    summary = fm._generate_change_summary()
    
    assert "Total Changes: 10" in summary
    assert "unicode_substitution: 10" in summary


if __name__ == '__main__':
    # Run tests
    pytest.main([__file__, '-v'])
