#!/usr/bin/env python3
"""
Buat dokumen DOCX demo untuk testing sistem turnitin bypass
"""
import docx
from pathlib import Path

def create_demo_docx():
    doc = docx.Document()
    
    # Header
    doc.add_heading('IMPLEMENTASI SISTEM E-COMMERCE', 0)
    doc.add_heading('BAB I PENDAHULUAN', level=1) 
    
    # Latar Belakang
    doc.add_heading('A. Latar Belakang', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('Era digital saat ini telah mengubah cara konsumen berbelanja. ')
    p1.add_run('Platform e-commerce seperti Shopee, Tokopedia, dan Lazada berkembang pesat ')
    p1.add_run('dengan jumlah pengunjung mencapai jutaan setiap harinya. ')
    p1.add_run('Fenomena ini menunjukkan pergeseran perilaku konsumen dari belanja konvensional ke belanja online.')
    
    p2 = doc.add_paragraph()
    p2.add_run('Keputusan pembelian konsumen dipengaruhi oleh berbagai faktor seperti kualitas produk, ')
    p2.add_run('harga, kemudahan akses, dan kepercayaan terhadap platform. ')
    p2.add_run('Menurut Kotler & Keller (2016), proses keputusan pembelian merupakan hasil dari ')
    p2.add_run('langkah-langkah yang dilakukan konsumen mulai dari pengenalan kebutuhan hingga evaluasi pasca pembelian.')
    
    # Rumusan Masalah  
    doc.add_heading('B. Rumusan Masalah', level=2)
    p3 = doc.add_paragraph('Berdasarkan latar belakang yang telah diuraikan, maka rumusan masalah dalam penelitian ini adalah sebagai berikut:')
    doc.add_paragraph('1. Bagaimana pengaruh kualitas produk terhadap keputusan pembelian online?')
    doc.add_paragraph('2. Faktor apa saja yang mempengaruhi kepercayaan konsumen pada platform e-commerce?')
    
    # BAB II
    doc.add_heading('BAB II TINJAUAN PUSTAKA', level=1)
    doc.add_heading('A. Landasan Teori', level=2)
    
    p4 = doc.add_paragraph()
    p4.add_run('E-commerce atau electronic commerce adalah proses jual beli produk secara elektronik ')
    p4.add_run('melalui jaringan komputer dan internet. Platform e-commerce menyediakan berbagai ')
    p4.add_run('kategori produk mulai dari fashion, elektronik, hingga kebutuhan rumah tangga.')
    
    # BAB III
    doc.add_heading('BAB III METODE PENELITIAN', level=1)
    doc.add_heading('A. Tempat dan Waktu Penelitian', level=2)
    
    p5 = doc.add_paragraph()
    p5.add_run('Penelitian ini dilakukan di Fakultas Teknik Informatika dengan jangka waktu 2 bulan. ')
    p5.add_run('Pemilihan lokasi penelitian berdasarkan pertimbangan kemudahan akses dan kelengkapan fasilitas.')
    
    # Save
    output_path = Path('input/demo_document.docx')
    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
    print(f"✅ Demo DOCX created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_demo_docx()