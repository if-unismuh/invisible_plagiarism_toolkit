#!/usr/bin/env python3
"""
Invisible Plagiarism Toolkit - Main CLI
Professional Document Manipulation System

Alur Sistem:
1. Upload 2 dokumen: original.docx + turnitin_report.pdf
2. OCR PDF dengan ocrmypdf untuk text extraction
3. Analisis highlight/flag dari PDF
4. Aplikasi manipulasi sesuai teknik:
   🔤 Unicode Steganography
   👻 Invisible Characters
   📑 Header Manipulation
   📋 Metadata Manipulation
   🔍 Verification System
   📊 Comprehensive Reporting
   🎯 Multiple Processing Modes

Author: DevNoLife
Version: 1.0.0
"""

import sys
import os
import argparse
import json
import subprocess
from pathlib import Path
from typing import Optional, List, Dict, Any

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from core.invisible_manipulator import InvisibleManipulator
from core.risk_analyzer import RiskAnalyzer
from extractors.pdf_colored_ocr_extractor import extract_colored_regions
from processors.flagged_selection_builder import build_selection, load_segments
from processors.targeted_text_matcher import TargetedTextMatcher, match_highlights_to_docx
from processors.header_protector import HeaderProtector
# from processors.targeted_invisible_applier import load_selection as load_flagged_selection
from utils.logger_config import setup_logger
from utils.performance_monitor import PerformanceMonitor

class PlagiarismBypassCLI:
    def __init__(self, workspace_dir: str = "workspace"):
        self.workspace = Path(workspace_dir)
        self.logger = setup_logger(__name__)
        self.monitor = PerformanceMonitor()
        self.setup_workspace()
        
    def setup_workspace(self):
        """Setup workspace directories"""
        directories = [
            "input/original",
            "input/turnitin", 
            "output/processed",
            "output/analysis",
            "output/reports",
            "temp"
        ]
        
        for dir_path in directories:
            (self.workspace / dir_path).mkdir(parents=True, exist_ok=True)
            
        self.logger.info(f"Workspace initialized at: {self.workspace.absolute()}")
        
    def check_dependencies(self) -> bool:
        """Check if required dependencies are installed"""
        try:
            # Check ocrmypdf
            result = subprocess.run(['ocrmypdf', '--version'], 
                                  capture_output=True, text=True)
            if result.returncode != 0:
                self.logger.error("ocrmypdf not found. Install with: sudo apt install ocrmypdf")
                return False
                
            self.logger.info(f"Dependencies OK - {result.stdout.strip()}")
            return True
            
        except FileNotFoundError:
            self.logger.error("ocrmypdf not found. Install with: sudo apt install ocrmypdf")
            return False
    
    def find_input_files(self) -> tuple[Optional[Path], Optional[Path]]:
        """Find original DOCX and Turnitin PDF files"""
        original_dir = self.workspace / "input" / "original"
        turnitin_dir = self.workspace / "input" / "turnitin"
        
        # Find DOCX files
        docx_files = list(original_dir.glob("*.docx"))
        original_file = docx_files[0] if docx_files else None
        
        # Find PDF files
        pdf_files = list(turnitin_dir.glob("*.pdf"))
        turnitin_file = pdf_files[0] if pdf_files else None
        
        if original_file:
            self.logger.info(f"Found original document: {original_file.name}")
        else:
            self.logger.warning("No DOCX file found in workspace/input/original/")
            
        if turnitin_file:
            self.logger.info(f"Found Turnitin report: {turnitin_file.name}")  
        else:
            self.logger.warning("No PDF file found in workspace/input/turnitin/")
            
        return original_file, turnitin_file
    
    def ocr_pdf(self, input_pdf: Path, output_pdf: Path, force: bool = True) -> bool:
        """Convert PDF to text-searchable using ocrmypdf"""
        self.logger.info(f"Converting PDF to text-searchable: {input_pdf.name}")
        
        cmd = ['ocrmypdf', str(input_pdf), str(output_pdf)]
        if force:
            cmd.append('--force-ocr')
        cmd.extend(['--language', 'ind+eng', '--optimize', '1'])
        
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            if result.returncode == 0:
                self.logger.info(f"OCR completed successfully: {output_pdf.name}")
                return True
            else:
                self.logger.error(f"OCR failed: {result.stderr}")
                return False
                
        except subprocess.TimeoutExpired:
            self.logger.error("OCR timeout after 5 minutes")
            return False
        except Exception as e:
            self.logger.error(f"OCR error: {e}")
            return False
    
    def extract_highlights(self, pdf_path: Path, mode: str = "balanced") -> List[Dict[str, Any]]:
        """Extract colored highlights from PDF"""
        self.logger.info(f"Extracting highlights with {mode} mode")
        
        # Mode settings
        settings = {
            "stealth": {"min_area": 1500, "aggressive": False, "max_coverage": 0.30},
            "balanced": {"min_area": 1200, "aggressive": True, "max_coverage": 0.50}, 
            "aggressive": {"min_area": 800, "aggressive": True, "max_coverage": 0.70}
        }
        
        config = settings.get(mode, settings["balanced"])
        
        try:
            highlights = extract_colored_regions(
                pdf_path,
                min_area=config["min_area"],
                aggressive=config["aggressive"],
                max_coverage=config["max_coverage"],
                merge=True,
                ocr_lang="ind+eng"
            )
            
            self.logger.info(f"Extracted {len(highlights)} highlighted segments")
            return highlights
            
        except Exception as e:
            self.logger.error(f"Highlight extraction failed: {e}")
            return []
    
    def filter_priority_highlights(self, highlights: List[Dict[str, Any]], 
                                 mode: str = "balanced") -> List[Dict[str, Any]]:
        """Filter highlights by Turnitin color priorities"""
        
        # Turnitin color priorities
        priority_colors = {
            "high": ["red", "green", "blue", "magenta"],
            "medium": ["orange", "cyan", "yellow"],
            "low": ["pink", "gray", "light"]
        }

        # Mode filtering
        if mode == "stealth":
            include_colors = set(priority_colors["high"])
            min_length = 15
            min_confidence = 0.45
            max_distance = 65.0
        elif mode == "balanced":
            include_colors = set(priority_colors["high"] + priority_colors["medium"])
            min_length = 10
            min_confidence = 0.35
            max_distance = 75.0
        else:  # aggressive
            include_colors = set(priority_colors["high"] + priority_colors["medium"] + priority_colors["low"])
            min_length = 6
            min_confidence = 0.25
            max_distance = 85.0

        # Build selection
        selection = build_selection(
            highlights, 
            min_length=min_length,
            include=include_colors,
            exclude=set(),
            dedupe=True,
            min_confidence=min_confidence,
            max_color_distance=max_distance
        )

        self.logger.info(f"Filtered to {len(selection)} priority segments ({mode} mode)")
        return selection
    
    def apply_manipulations(self, original_docx: Path, highlights: List[Dict[str, Any]],
                          mode: str = "balanced", use_targeted: bool = True) -> Optional[Path]:
        """
        Apply invisible manipulations to original document

        Args:
            original_docx: Path to original DOCX file
            highlights: List of highlights extracted from Turnitin PDF
            mode: Processing mode (stealth/balanced/aggressive)
            use_targeted: If True, only modify paragraphs that match highlights

        Returns:
            Path to processed document or None if failed
        """

        # Output path
        output_path = self.workspace / "output" / "processed" / f"{original_docx.stem}_{mode}_processed.docx"

        try:
            if use_targeted and highlights:
                self.logger.info("Using TARGETED manipulation (only modifying flagged paragraphs)")

                # Step 1: Match highlights to exact paragraphs in DOCX
                self.logger.info("Matching highlights to document paragraphs...")
                matches = match_highlights_to_docx(
                    str(original_docx),
                    highlights,
                    min_similarity=0.80  # 80% similarity threshold
                )

                if not matches:
                    self.logger.warning("No paragraph matches found! Falling back to full document processing")
                    use_targeted = False
                else:
                    self.logger.info(f"Found {len(matches)} paragraph matches to modify")

                    # Step 2: Apply targeted manipulations
                    return self._apply_targeted_manipulations(
                        original_docx, matches, output_path, mode
                    )

            if not use_targeted:
                # Fallback: Full document processing (original behavior)
                self.logger.info("Using FULL document manipulation")
                manipulator = InvisibleManipulator(verbose=True)

                result = manipulator.apply_invisible_manipulation(
                    str(original_docx),
                    str(output_path),
                    dry_run=False
                )

                if result and result.get('output_file'):
                    self.logger.info(f"Document processed successfully: {output_path.name}")
                    return Path(result['output_file'])
                else:
                    self.logger.error("Document processing failed")
                    return None

        except Exception as e:
            self.logger.error(f"Manipulation failed: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return None

    def _apply_targeted_manipulations(self, original_docx: Path, matches: List[Dict[str, Any]],
                                     output_path: Path, mode: str) -> Optional[Path]:
        """Apply manipulations only to matched paragraphs"""
        import docx
        from core.unicode_steganography import UnicodeSteg

        try:
            # Load document
            doc = docx.Document(str(original_docx))
            steg = UnicodeSteg()

            # Mode settings
            rates = {
                "stealth": {"unicode": 0.02, "zero_width": 0.03, "max_changes": 2},
                "balanced": {"unicode": 0.04, "zero_width": 0.06, "max_changes": 3},
                "aggressive": {"unicode": 0.07, "zero_width": 0.10, "max_changes": 5}
            }
            config = rates.get(mode, rates["balanced"])

            modified_count = 0
            total_changes = 0

            # STEP 1: Protect ALL standard headers first (BAB I, PENDAHULUAN, etc.)
            self.logger.info("🛡️  Protecting standard headers...")
            protector = HeaderProtector()
            header_stats = protector.protect_all_headers(doc, aggressiveness=mode)

            self.logger.info(f"   - Headers found: {header_stats['total_headers_found']}")
            self.logger.info(f"   - Headers protected: {header_stats['headers_protected']}")
            self.logger.info(f"   - Header changes: {header_stats['total_changes']}")

            # STEP 2: Process each matched paragraph (flagged by Turnitin)
            for match in matches:
                para_idx = match['paragraph_index']
                paragraph = doc.paragraphs[para_idx]
                original_text = paragraph.text

                if not original_text.strip():
                    continue

                self.logger.debug(f"Processing paragraph {para_idx}: {len(original_text)} chars")

                # Apply Unicode substitution
                modified_text, stats = steg.apply_strategic_substitution(
                    original_text,
                    aggressiveness=config['unicode']
                )
                sub_count = stats.get('total_changes', 0)

                self.logger.debug(f"Unicode substitutions: {sub_count}")

                # Apply invisible characters
                modified_text_with_invisible = self._insert_invisible_chars(
                    modified_text,
                    rate=config['zero_width']
                )

                # Update paragraph runs (preserve formatting)
                if modified_text_with_invisible != original_text:
                    # Clear existing runs
                    for run in paragraph.runs:
                        run.text = ''

                    # Add new text as single run
                    if paragraph.runs:
                        paragraph.runs[0].text = modified_text_with_invisible
                    else:
                        paragraph.add_run(modified_text_with_invisible)

                    modified_count += 1
                    total_changes += sub_count

                    self.logger.debug(f"✓ Modified paragraph {para_idx}: {sub_count} changes")
                else:
                    self.logger.debug(f"✗ No changes for paragraph {para_idx}")

            # Save document
            doc.save(str(output_path))

            self.logger.info(f"✅ Targeted manipulation complete:")
            self.logger.info(f"   - Headers protected: {header_stats['headers_protected']}")
            self.logger.info(f"   - Flagged paragraphs modified: {modified_count}/{len(matches)}")
            self.logger.info(f"   - Total changes: {total_changes + header_stats['total_changes']}")
            self.logger.info(f"   - Output: {output_path.name}")

            return output_path

        except Exception as e:
            self.logger.error(f"Targeted manipulation failed: {e}")
            import traceback
            self.logger.error(traceback.format_exc())
            return None

    def _insert_invisible_chars(self, text: str, rate: float = 0.05) -> str:
        """Insert zero-width characters into text"""
        import random

        zero_width_chars = ['\u200B', '\u200C', '\u200D', '\uFEFF']
        words = text.split()
        modified_words = []

        for word in words:
            if random.random() < rate and len(word) > 3:
                # Insert zero-width char in middle of word
                mid = len(word) // 2
                char = random.choice(zero_width_chars)
                word = word[:mid] + char + word[mid:]

            modified_words.append(word)

        return ' '.join(modified_words)
    
    def generate_report(self, original_file: Path, turnitin_file: Path, 
                       processed_file: Optional[Path], highlights: List[Dict[str, Any]],
                       selections: List[Dict[str, Any]], mode: str) -> Path:
        """Generate comprehensive processing report"""
        
        report_data = {
            "processing_info": {
                "mode": mode,
                "timestamp": self.monitor.get_current_time(),
                "processing_time": self.monitor.get_elapsed_time()
            },
            "input_files": {
                "original_document": str(original_file.name),
                "turnitin_report": str(turnitin_file.name)
            },
            "analysis_results": {
                "total_highlights": len(highlights),
                "selected_segments": len(selections),
                "color_distribution": self._get_color_distribution(highlights)
            },
            "output_files": {
                "processed_document": str(processed_file.name) if processed_file else None
            },
            "techniques_applied": [
                "🔤 Unicode Steganography",
                "👻 Invisible Characters", 
                "📑 Header Manipulation",
                "📋 Metadata Manipulation"
            ]
        }
        
        report_path = self.workspace / "output" / "reports" / f"processing_report_{mode}.json"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)
            
        self.logger.info(f"Report generated: {report_path.name}")
        return report_path
    
    def _get_color_distribution(self, highlights: List[Dict[str, Any]]) -> Dict[str, int]:
        """Calculate color distribution from highlights"""
        from collections import Counter
        colors = [h.get('color', 'unknown') for h in highlights]
        return dict(Counter(colors))
    
    def process_documents(self, mode: str = "balanced") -> bool:
        """Main processing pipeline"""
        self.monitor.start()
        
        self.logger.info("=" * 60)
        self.logger.info("🚀 INVISIBLE PLAGIARISM TOOLKIT - PROCESSING STARTED")
        self.logger.info("=" * 60)
        
        # Check dependencies
        if not self.check_dependencies():
            return False
            
        # Find input files
        original_file, turnitin_file = self.find_input_files()
        
        if not original_file or not turnitin_file:
            self.logger.error("Required input files not found!")
            self.logger.info("Please place files in:")
            self.logger.info(f"  - Original DOCX: {self.workspace / 'input' / 'original'}")
            self.logger.info(f"  - Turnitin PDF: {self.workspace / 'input' / 'turnitin'}")
            return False
        
        # OCR PDF for text extraction
        ocr_pdf_path = self.workspace / "temp" / f"ocr_{turnitin_file.name}"
        if not self.ocr_pdf(turnitin_file, ocr_pdf_path):
            self.logger.warning("OCR failed, using original PDF")
            ocr_pdf_path = turnitin_file
        
        # Extract highlights
        highlights = self.extract_highlights(ocr_pdf_path, mode)
        if not highlights:
            self.logger.error("No highlights found in PDF")
            return False
        
        # Filter priority highlights
        selections = self.filter_priority_highlights(highlights, mode)
        if not selections:
            self.logger.error("No priority highlights selected")
            return False
        
        # Apply manipulations (now using TARGETED approach)
        # Pass highlights instead of selections for better paragraph matching
        processed_file = self.apply_manipulations(
            original_file,
            highlights,  # Use all highlights for matching
            mode,
            use_targeted=True  # Enable targeted modification
        )
        if not processed_file:
            self.logger.error("Document manipulation failed")
            return False
        
        # Generate report
        report_path = self.generate_report(
            original_file, turnitin_file, processed_file,
            highlights, selections, mode
        )

        # Analyze detection risk
        self.logger.info("Analyzing detection risk...")
        risk_analyzer = RiskAnalyzer()

        # Load changes log if exists
        changes_log_path = processed_file.with_suffix('.changes.json')
        changes_log = None
        if changes_log_path.exists():
            with open(changes_log_path, 'r', encoding='utf-8') as f:
                changes_log = json.load(f)

        risk_score = risk_analyzer.analyze_document(str(processed_file), changes_log)
        risk_report = risk_analyzer.generate_report(
            risk_score,
            output_path=str(self.workspace / "output" / "reports" / f"risk_analysis_{mode}.txt")
        )

        # Final summary
        self.logger.info("=" * 60)
        self.logger.info("✅ PROCESSING COMPLETED SUCCESSFULLY")
        self.logger.info("=" * 60)
        self.logger.info(f"📄 Processed Document: {processed_file}")
        self.logger.info(f"📊 Report: {report_path}")
        self.logger.info(f"🔍 Risk Analysis: {risk_score.risk_level} ({risk_score.overall_score:.1f}/100)")
        self.logger.info(f"⏱️  Processing Time: {self.monitor.get_elapsed_time():.2f}s")

        # Print risk report
        print("\n" + risk_report)

        return True

def create_cli_parser() -> argparse.ArgumentParser:
    """Create CLI argument parser"""
    parser = argparse.ArgumentParser(
        description="Invisible Plagiarism Toolkit - Professional Document Manipulation",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Processing Modes:
  stealth    - Minimal modifications, highest invisibility
  balanced   - Moderate modifications, good invisibility/coverage balance  
  aggressive - Maximum modifications, highest bypass potential

Usage Examples:
  python main.py --mode balanced
  python main.py --mode stealth --workspace /custom/path
  python main.py --help

Before running:
  1. Place original DOCX in: workspace/input/original/
  2. Place Turnitin PDF in: workspace/input/turnitin/
        """
    )
    
    parser.add_argument(
        '--mode', 
        choices=['stealth', 'balanced', 'aggressive'],
        default='balanced',
        help='Processing mode (default: balanced)'
    )
    
    parser.add_argument(
        '--workspace',
        default='workspace',
        help='Workspace directory path (default: workspace)'
    )
    
    parser.add_argument(
        '--check-deps',
        action='store_true',
        help='Check dependencies only'
    )
    
    parser.add_argument(
        '--debug',
        action='store_true',
        help='Enable debug mode with visual flags (highlights changes in document)'
    )
    
    parser.add_argument(
        '--no-change-log',
        action='store_true',
        help='Disable embedded change log in output document'
    )

    parser.add_argument(
        '--analyze-risk',
        metavar='FILE',
        help='Analyze detection risk for an existing document'
    )

    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Preview changes without modifying files (dry-run mode)'
    )

    return parser

def main():
    """Main CLI entry point"""
    parser = create_cli_parser()
    args = parser.parse_args()
    
    # Initialize CLI
    cli = PlagiarismBypassCLI(args.workspace)
    
    # Check dependencies only
    if args.check_deps:
        if cli.check_dependencies():
            print("✅ All dependencies are properly installed")
            return 0
        else:
            print("❌ Missing dependencies")
            return 1

    # Analyze risk for existing document
    if args.analyze_risk:
        doc_path = Path(args.analyze_risk)
        if not doc_path.exists():
            print(f"❌ File not found: {doc_path}")
            return 1

        print(f"🔍 Analyzing detection risk for: {doc_path.name}")

        # Look for changes log
        changes_log_path = doc_path.with_suffix('.changes.json')
        changes_log = None
        if changes_log_path.exists():
            with open(changes_log_path, 'r', encoding='utf-8') as f:
                changes_log = json.load(f)
            print(f"📊 Found change log: {changes_log_path.name}")

        analyzer = RiskAnalyzer()
        risk_score = analyzer.analyze_document(str(doc_path), changes_log)
        risk_report = analyzer.generate_report(risk_score)

        print("\n" + risk_report)
        return 0

    # Update config for debug mode if enabled
    if args.debug:
        print("🐛 Debug mode enabled - visual flags will be added to document")
        config_path = Path("config.json")
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
            config.setdefault('debug', {})['enable_visual_flags'] = True
            with open(config_path, 'w') as f:
                json.dump(config, f, indent=2)

    if args.no_change_log:
        print("📝 Change log disabled")
        config_path = Path("config.json")
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
            config.setdefault('debug', {})['enable_change_log'] = False
            config['debug']['export_change_log'] = False
            with open(config_path, 'w') as f:
                json.dump(config, f, indent=2)
    
    # Process documents
    success = cli.process_documents(args.mode)
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())
