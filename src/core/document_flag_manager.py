#!/usr/bin/env python3
"""
Document Flag Manager
====================
Menambahkan visual flags dan markers di dokumen DOCX untuk tracking perubahan.
Berguna untuk debugging, verification, dan audit trail.

Features:
- Comment annotations pada teks yang dimodifikasi
- Hidden text markers (track changes tanpa visible)
- Highlight colors untuk debug mode
- Custom document properties untuk metadata tracking
- Change log embedded dalam dokumen

Author: DevNoLife
Version: 1.0
"""

import logging
from datetime import datetime
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class ChangeType:
    """Constants untuk tipe perubahan"""
    UNICODE_SUBSTITUTION = "unicode_substitution"
    ZERO_WIDTH_INSERTION = "zero_width_insertion"
    HEADER_MODIFICATION = "header_modification"
    METADATA_CHANGE = "metadata_change"
    PARAPHRASE = "paraphrase"


class DocumentFlagManager:
    """Manage flags dan markers untuk tracking perubahan dokumen"""
    
    def __init__(self, enable_visual_flags: bool = False, enable_comments: bool = True):
        """
        Initialize Flag Manager
        
        Args:
            enable_visual_flags: Aktifkan highlight visual (untuk debug)
            enable_comments: Aktifkan comment annotations
        """
        self.enable_visual_flags = enable_visual_flags
        self.enable_comments = enable_comments
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # Track all changes
        self.changes: List[Dict[str, Any]] = []
        
        # Color mapping untuk visual flags
        self.color_map = {
            ChangeType.UNICODE_SUBSTITUTION: WD_COLOR_INDEX.YELLOW,
            ChangeType.ZERO_WIDTH_INSERTION: WD_COLOR_INDEX.BRIGHT_GREEN,
            ChangeType.HEADER_MODIFICATION: WD_COLOR_INDEX.TURQUOISE,
            ChangeType.METADATA_CHANGE: WD_COLOR_INDEX.GRAY_25,
            ChangeType.PARAPHRASE: WD_COLOR_INDEX.PINK,
        }
        
        self.logger.info("DocumentFlagManager initialized (visual=%s, comments=%s)", 
                        enable_visual_flags, enable_comments)
    
    def add_change_record(self, change_type: str, location: str, 
                         original: str, modified: str, details: Optional[Dict] = None):
        """
        Record a change for tracking
        
        Args:
            change_type: Type of modification (use ChangeType constants)
            location: Location identifier (e.g., "paragraph_5", "header_2")
            original: Original text
            modified: Modified text
            details: Additional details dict
        """
        change_record = {
            'timestamp': datetime.now().isoformat(),
            'type': change_type,
            'location': location,
            'original': original[:100],  # Limit length
            'modified': modified[:100],
            'details': details or {}
        }
        
        self.changes.append(change_record)
        self.logger.debug("Change recorded: %s at %s", change_type, location)
    
    def flag_paragraph(self, paragraph, change_type: str, message: str = None):
        """
        Add flag to a paragraph
        
        Args:
            paragraph: python-docx Paragraph object
            change_type: Type of modification
            message: Optional custom message
        """
        try:
            # Add comment if enabled
            if self.enable_comments:
                self._add_comment_to_paragraph(paragraph, change_type, message)
            
            # Add visual highlight if enabled (debug mode)
            if self.enable_visual_flags:
                self._add_highlight_to_paragraph(paragraph, change_type)
            
            # Add hidden tracking marker
            self._add_hidden_marker(paragraph, change_type)
            
        except Exception as e:
            self.logger.warning("Failed to flag paragraph: %s", e)
    
    def _add_comment_to_paragraph(self, paragraph, change_type: str, message: str = None):
        """Add comment annotation to paragraph"""
        try:
            # Create comment text
            comment_text = message or f"Modified: {change_type}"
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            full_comment = f"[IPT] {comment_text}\nType: {change_type}\nTime: {timestamp}"
            
            # Note: python-docx doesn't have direct comment API
            # We'll add it as a custom XML property instead
            # For full comment support, would need direct XML manipulation
            
            self.logger.debug("Comment prepared for paragraph: %s", comment_text)
            
        except Exception as e:
            self.logger.warning("Failed to add comment: %s", e)
    
    def _add_highlight_to_paragraph(self, paragraph, change_type: str):
        """Add visual highlight to paragraph (debug mode only)"""
        try:
            color = self.color_map.get(change_type, WD_COLOR_INDEX.YELLOW)
            
            for run in paragraph.runs:
                run.font.highlight_color = color
            
            self.logger.debug("Highlight added to paragraph: %s", change_type)
            
        except Exception as e:
            self.logger.warning("Failed to add highlight: %s", e)
    
    def _add_hidden_marker(self, paragraph, change_type: str):
        """Add hidden marker to paragraph for tracking"""
        try:
            # Add a zero-width marker at the end with special pattern
            # Format: \u200B\u200C\u200D (specific sequence for IPT marker)
            marker = f"\u200B\u200C\u200D[{change_type}]"
            
            if paragraph.runs:
                # Add to last run
                last_run = paragraph.runs[-1]
                last_run.text += marker
            else:
                # Create new run if none exist
                paragraph.add_run(marker)
            
            self.logger.debug("Hidden marker added: %s", change_type)
            
        except Exception as e:
            self.logger.warning("Failed to add hidden marker: %s", e)
    
    def flag_run(self, run, change_type: str, original_text: str, modified_text: str):
        """
        Flag a specific run (inline text) that was modified
        
        Args:
            run: python-docx Run object
            change_type: Type of modification
            original_text: Original text
            modified_text: Modified text
        """
        try:
            # Add visual highlight if enabled
            if self.enable_visual_flags:
                color = self.color_map.get(change_type, WD_COLOR_INDEX.YELLOW)
                run.font.highlight_color = color
            
            # Add hidden marker
            marker = f"\u200B[{change_type}]\u200C"
            run.text = modified_text + marker
            
            self.logger.debug("Run flagged: %s (%s chars)", change_type, len(modified_text))
            
        except Exception as e:
            self.logger.warning("Failed to flag run: %s", e)
    
    def add_document_properties(self, doc: Document):
        """
        Add custom document properties for tracking
        
        Args:
            doc: python-docx Document object
        """
        try:
            core_props = doc.core_properties
            
            # Add custom properties
            core_props.comments = self._generate_change_summary()
            core_props.subject = "IPT Modified Document"
            core_props.keywords = f"IPT, Modified, {datetime.now().strftime('%Y-%m-%d')}"
            
            # Custom property for change count
            # Note: python-docx doesn't support custom properties directly
            # Would need XML manipulation for full support
            
            self.logger.info("Document properties updated with tracking info")
            
        except Exception as e:
            self.logger.warning("Failed to add document properties: %s", e)
    
    def _generate_change_summary(self) -> str:
        """Generate a summary of all changes"""
        summary_lines = [
            "=== IPT Modification Summary ===",
            f"Total Changes: {len(self.changes)}",
            f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            ""
        ]
        
        # Count by type
        type_counts = {}
        for change in self.changes:
            change_type = change['type']
            type_counts[change_type] = type_counts.get(change_type, 0) + 1
        
        summary_lines.append("Changes by Type:")
        for change_type, count in type_counts.items():
            summary_lines.append(f"  - {change_type}: {count}")
        
        return "\n".join(summary_lines)
    
    def add_change_log_section(self, doc: Document):
        """
        Add a hidden change log section at the end of document
        
        Args:
            doc: python-docx Document object
        """
        try:
            # Add page break
            doc.add_page_break()
            
            # Add title
            title = doc.add_paragraph()
            title_run = title.add_run("IPT MODIFICATION LOG")
            title_run.bold = True
            title_run.font.size = 14
            
            # Make it hidden text
            title_run.font.hidden = True
            
            # Add summary
            summary = doc.add_paragraph()
            summary_run = summary.add_run(self._generate_change_summary())
            summary_run.font.hidden = True
            summary_run.font.size = 10
            
            # Add detailed changes
            if self.changes:
                details_para = doc.add_paragraph()
                details_run = details_para.add_run("\n=== Detailed Changes ===\n\n")
                details_run.font.hidden = True
                
                for i, change in enumerate(self.changes[:50], 1):  # Limit to 50
                    change_text = (
                        f"{i}. {change['type']} at {change['location']}\n"
                        f"   Time: {change['timestamp']}\n"
                        f"   Original: {change['original'][:50]}...\n"
                        f"   Modified: {change['modified'][:50]}...\n\n"
                    )
                    change_run = details_para.add_run(change_text)
                    change_run.font.hidden = True
                    change_run.font.size = 9
            
            self.logger.info("Change log section added to document (hidden)")
            
        except Exception as e:
            self.logger.warning("Failed to add change log section: %s", e)
    
    def detect_previous_modifications(self, doc: Document) -> bool:
        """
        Detect if document was previously modified by IPT
        
        Args:
            doc: python-docx Document object
            
        Returns:
            bool: True if document has IPT markers
        """
        try:
            # Check document properties
            if hasattr(doc.core_properties, 'subject'):
                if 'IPT Modified' in str(doc.core_properties.subject):
                    self.logger.warning("Document appears to be already modified by IPT!")
                    return True
            
            # Check for hidden markers in text
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if '\u200B\u200C\u200D[' in text:  # IPT marker pattern
                    self.logger.warning("Found IPT markers in document!")
                    return True
            
            return False
            
        except Exception as e:
            self.logger.warning("Error detecting previous modifications: %s", e)
            return False
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        Get statistics about tracked changes
        
        Returns:
            dict: Statistics dictionary
        """
        stats = {
            'total_changes': len(self.changes),
            'changes_by_type': {},
            'first_change': None,
            'last_change': None
        }
        
        if self.changes:
            stats['first_change'] = self.changes[0]['timestamp']
            stats['last_change'] = self.changes[-1]['timestamp']
            
            # Count by type
            for change in self.changes:
                change_type = change['type']
                stats['changes_by_type'][change_type] = \
                    stats['changes_by_type'].get(change_type, 0) + 1
        
        return stats
    
    def export_change_log(self, output_path: str):
        """
        Export change log to JSON file
        
        Args:
            output_path: Path to save JSON file
        """
        import json
        
        try:
            log_data = {
                'metadata': {
                    'timestamp': datetime.now().isoformat(),
                    'total_changes': len(self.changes),
                    'visual_flags_enabled': self.enable_visual_flags,
                    'comments_enabled': self.enable_comments
                },
                'statistics': self.get_statistics(),
                'changes': self.changes
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(log_data, f, indent=2, ensure_ascii=False)
            
            self.logger.info("Change log exported to: %s", output_path)
            
        except Exception as e:
            self.logger.error("Failed to export change log: %s", e)


class FlaggedTextBuilder:
    """Helper class untuk membangun teks dengan flags"""
    
    def __init__(self, flag_manager: DocumentFlagManager):
        self.flag_manager = flag_manager
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def create_flagged_run(self, paragraph, text: str, change_type: str):
        """
        Create a new run with flags
        
        Args:
            paragraph: Paragraph object
            text: Text content
            change_type: Type of modification
            
        Returns:
            Run object
        """
        run = paragraph.add_run(text)
        
        if self.flag_manager.enable_visual_flags:
            color = self.flag_manager.color_map.get(change_type, WD_COLOR_INDEX.YELLOW)
            run.font.highlight_color = color
        
        # Add hidden marker
        marker = f"\u200B[{change_type}]\u200C"
        run.text += marker
        
        return run
    
    def replace_text_with_flag(self, paragraph, old_text: str, new_text: str, 
                              change_type: str) -> bool:
        """
        Replace text in paragraph and add flag
        
        Args:
            paragraph: Paragraph object
            old_text: Text to replace
            new_text: Replacement text
            change_type: Type of modification
            
        Returns:
            bool: Success status
        """
        try:
            # Simple replacement (more complex logic would preserve formatting)
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                self.flag_manager.flag_paragraph(paragraph, change_type)
                return True
            return False
            
        except Exception as e:
            self.logger.warning("Failed to replace text with flag: %s", e)
            return False


def create_flag_manager(config: Dict[str, Any]) -> DocumentFlagManager:
    """
    Factory function to create flag manager from config
    
    Args:
        config: Configuration dictionary
        
    Returns:
        DocumentFlagManager instance
    """
    # Check if debug mode is enabled
    debug_mode = config.get('debug', {}).get('enable_visual_flags', False)
    enable_comments = config.get('debug', {}).get('enable_comments', True)
    
    return DocumentFlagManager(
        enable_visual_flags=debug_mode,
        enable_comments=enable_comments
    )
