"""
Targeted Text Matcher
Match extracted highlights from Turnitin PDF to exact locations in DOCX

Author: DevNoLife
Version: 1.0
"""

import logging
import re
from typing import List, Dict, Any, Tuple, Optional
from difflib import SequenceMatcher
import docx
from docx.text.paragraph import Paragraph


class TargetedTextMatcher:
    """Match highlighted text from PDF to exact paragraphs in DOCX"""

    def __init__(self, min_similarity: float = 0.85):
        """
        Initialize matcher

        Args:
            min_similarity: Minimum similarity ratio for text matching (0.0-1.0)
        """
        self.logger = logging.getLogger(self.__class__.__name__)
        self.min_similarity = min_similarity

    def find_matches(self, doc: docx.Document, highlights: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Find exact paragraph matches for highlighted text

        Args:
            doc: DOCX document object
            highlights: List of highlight dicts from PDF extractor
                        Format: {text, color, page_number, bbox, ...}

        Returns:
            List of match dicts with format:
            {
                'paragraph_index': int,
                'paragraph': Paragraph object,
                'original_text': str,
                'highlight_text': str,
                'color': str,
                'similarity': float,
                'page_number': int,
                'match_type': 'exact' | 'partial' | 'fuzzy'
            }
        """
        self.logger.info(f"Matching {len(highlights)} highlights to document paragraphs...")

        matches = []
        all_text = self._extract_all_text(doc)

        for highlight in highlights:
            highlight_text = highlight.get('text', '').strip()
            if not highlight_text or len(highlight_text) < 10:
                continue  # Skip too short highlights

            # Try to find match
            match = self._find_best_match(doc, all_text, highlight)
            if match:
                matches.append(match)

        self.logger.info(f"Found {len(matches)} paragraph matches")
        return matches

    def _extract_all_text(self, doc: docx.Document) -> List[Tuple[int, str]]:
        """Extract all paragraph texts with their indices"""
        return [
            (i, para.text.strip())
            for i, para in enumerate(doc.paragraphs)
            if para.text.strip()
        ]

    def _find_best_match(self, doc: docx.Document, all_text: List[Tuple[int, str]],
                        highlight: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Find best matching paragraph for a highlight"""
        highlight_text = highlight.get('text', '').strip()
        highlight_clean = self._clean_text(highlight_text)

        best_match = None
        best_similarity = 0.0
        match_type = None

        for para_idx, para_text in all_text:
            para_clean = self._clean_text(para_text)

            # Try exact match first
            if highlight_clean == para_clean:
                best_match = para_idx
                best_similarity = 1.0
                match_type = 'exact'
                break

            # Try substring match (highlight might be part of paragraph)
            if highlight_clean in para_clean or para_clean in highlight_clean:
                similarity = self._calculate_similarity(highlight_clean, para_clean)
                if similarity > best_similarity and similarity >= self.min_similarity:
                    best_match = para_idx
                    best_similarity = similarity
                    match_type = 'partial'

            # Try fuzzy match
            similarity = self._calculate_similarity(highlight_clean, para_clean)
            if similarity > best_similarity and similarity >= self.min_similarity:
                best_match = para_idx
                best_similarity = similarity
                match_type = 'fuzzy'

        if best_match is not None:
            paragraph = doc.paragraphs[best_match]
            return {
                'paragraph_index': best_match,
                'paragraph': paragraph,
                'original_text': paragraph.text,
                'highlight_text': highlight_text,
                'color': highlight.get('color', 'unknown'),
                'similarity': best_similarity,
                'page_number': highlight.get('page_number', 0),
                'match_type': match_type,
                'bbox': highlight.get('bbox')
            }

        return None

    def _clean_text(self, text: str) -> str:
        """Clean text for comparison"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove punctuation variations
        text = text.replace('\n', ' ').replace('\r', ' ')
        # Lowercase for comparison
        text = text.lower().strip()
        return text

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity ratio between two texts"""
        return SequenceMatcher(None, text1, text2).ratio()

    def group_by_color(self, matches: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Group matches by Turnitin color priority"""
        grouped = {
            'high_priority': [],    # red, green, blue, magenta
            'medium_priority': [],  # orange, cyan, yellow
            'low_priority': []      # gray, pink, other
        }

        high_colors = {'red', 'green', 'blue', 'magenta'}
        medium_colors = {'orange', 'cyan', 'yellow'}

        for match in matches:
            color = match.get('color', '').lower()
            if color in high_colors:
                grouped['high_priority'].append(match)
            elif color in medium_colors:
                grouped['medium_priority'].append(match)
            else:
                grouped['low_priority'].append(match)

        return grouped

    def filter_duplicates(self, matches: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Remove duplicate matches (same paragraph matched multiple times)"""
        seen_paragraphs = set()
        unique_matches = []

        # Sort by similarity (highest first) to keep best matches
        sorted_matches = sorted(matches, key=lambda x: x['similarity'], reverse=True)

        for match in sorted_matches:
            para_idx = match['paragraph_index']
            if para_idx not in seen_paragraphs:
                seen_paragraphs.add(para_idx)
                unique_matches.append(match)

        self.logger.info(f"Filtered {len(matches)} matches to {len(unique_matches)} unique paragraphs")
        return unique_matches

    def get_statistics(self, matches: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get matching statistics"""
        if not matches:
            return {
                'total_matches': 0,
                'exact_matches': 0,
                'partial_matches': 0,
                'fuzzy_matches': 0,
                'by_color': {},
                'average_similarity': 0.0
            }

        match_types = {'exact': 0, 'partial': 0, 'fuzzy': 0}
        color_counts = {}
        total_similarity = 0.0

        for match in matches:
            match_type = match.get('match_type', 'fuzzy')
            match_types[match_type] = match_types.get(match_type, 0) + 1

            color = match.get('color', 'unknown')
            color_counts[color] = color_counts.get(color, 0) + 1

            total_similarity += match.get('similarity', 0.0)

        return {
            'total_matches': len(matches),
            'exact_matches': match_types['exact'],
            'partial_matches': match_types['partial'],
            'fuzzy_matches': match_types['fuzzy'],
            'by_color': color_counts,
            'average_similarity': total_similarity / len(matches) if matches else 0.0
        }


def match_highlights_to_docx(docx_path: str, highlights: List[Dict[str, Any]],
                             min_similarity: float = 0.85) -> List[Dict[str, Any]]:
    """
    Convenience function to match highlights to DOCX paragraphs

    Args:
        docx_path: Path to DOCX file
        highlights: List of highlights from PDF extractor
        min_similarity: Minimum similarity for fuzzy matching

    Returns:
        List of matched paragraphs with modification targets
    """
    doc = docx.Document(docx_path)
    matcher = TargetedTextMatcher(min_similarity=min_similarity)

    # Find all matches
    matches = matcher.find_matches(doc, highlights)

    # Remove duplicates (keep best match per paragraph)
    matches = matcher.filter_duplicates(matches)

    # Print statistics
    stats = matcher.get_statistics(matches)
    logger = logging.getLogger(__name__)
    logger.info(f"Matching Statistics: {stats}")

    return matches
