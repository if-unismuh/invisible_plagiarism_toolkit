"""
Header Protector
Protect standard headers from plagiarism detection

Standard headers seperti "BAB I PENDAHULUAN" WAJIB ada di skripsi
tapi Turnitin detect sebagai plagiarisme. Module ini apply invisible
techniques khusus untuk headers.

Author: DevNoLife
Version: 1.0
"""

import logging
import re
import random
from typing import List, Dict, Tuple, Optional
import docx
from docx.text.paragraph import Paragraph


class HeaderProtector:
    """Apply invisible techniques to standard academic headers"""

    # Standard Indonesian academic headers (WAJIB ada di skripsi)
    STANDARD_HEADERS = [
        # BAB headers
        r'BAB\s+I\b',
        r'BAB\s+II\b',
        r'BAB\s+III\b',
        r'BAB\s+IV\b',
        r'BAB\s+V\b',
        r'BAB\s+1\b',
        r'BAB\s+2\b',
        r'BAB\s+3\b',
        r'BAB\s+4\b',
        r'BAB\s+5\b',

        # Common chapter names
        r'\bPENDAHULUAN\b',
        r'\bTINJAUAN\s+PUSTAKA\b',
        r'\bMETODE\s+PENELITIAN\b',
        r'\bMETODOLOGI\s+PENELITIAN\b',
        r'\bHASIL\s+DAN\s+PEMBAHASAN\b',
        r'\bKESIMPULAN\s+DAN\s+SARAN\b',
        r'\bKESIMPULAN\b',
        r'\bSARAN\b',

        # Front matter
        r'\bABSTRAK\b',
        r'\bABSTRACT\b',
        r'\bKATA\s+PENGANTAR\b',
        r'\bDAFTAR\s+ISI\b',
        r'\bDAFTAR\s+TABEL\b',
        r'\bDAFTAR\s+GAMBAR\b',
        r'\bDAFTAR\s+PUSTAKA\b',
        r'\bLAMPIRAN\b',
    ]

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)

        # Unicode mappings (Latin → Cyrillic/Greek lookalikes)
        self.unicode_map = {
            # Uppercase
            'A': 'А',  # Cyrillic A
            'B': 'В',  # Cyrillic B
            'C': 'С',  # Cyrillic C
            'E': 'Е',  # Cyrillic E
            'H': 'Н',  # Cyrillic H
            'I': 'І',  # Cyrillic I
            'K': 'К',  # Cyrillic K
            'M': 'М',  # Cyrillic M
            'N': 'Н',  # Cyrillic N
            'O': 'О',  # Cyrillic O
            'P': 'Р',  # Cyrillic P
            'S': 'Ѕ',  # Cyrillic S
            'T': 'Т',  # Cyrillic T
            'X': 'Х',  # Cyrillic X

            # Lowercase
            'a': 'а',  # Cyrillic a
            'c': 'с',  # Cyrillic c
            'e': 'е',  # Cyrillic e
            'i': 'і',  # Cyrillic i
            'o': 'о',  # Cyrillic o
            'p': 'р',  # Cyrillic p
            's': 'ѕ',  # Cyrillic s
            'x': 'х',  # Cyrillic x
        }

        # Zero-width characters
        self.zero_width_chars = [
            '\u200B',  # Zero Width Space
            '\u200C',  # Zero Width Non-Joiner
            '\u200D',  # Zero Width Joiner
            '\uFEFF',  # Zero Width No-Break Space
        ]

    def is_standard_header(self, text: str) -> bool:
        """Check if text is a standard academic header"""
        text_upper = text.upper().strip()

        for pattern in self.STANDARD_HEADERS:
            if re.search(pattern, text_upper):
                return True

        return False

    def protect_header(self, text: str, aggressiveness: str = 'balanced') -> Tuple[str, Dict]:
        """
        Apply invisible protection to header text

        Args:
            text: Header text (e.g., "BAB I PENDAHULUAN")
            aggressiveness: 'stealth', 'balanced', or 'aggressive'

        Returns:
            (modified_text, stats)
        """
        if not text.strip():
            return text, {'changes': 0, 'techniques': []}

        stats = {
            'original': text,
            'changes': 0,
            'techniques': []
        }

        # Determine rates based on aggressiveness
        rates = {
            'stealth': {
                'unicode': 0.3,     # 30% chars
                'zero_width': 0.2,  # 20% spaces
            },
            'balanced': {
                'unicode': 0.5,     # 50% chars
                'zero_width': 0.4,  # 40% spaces
            },
            'aggressive': {
                'unicode': 0.7,     # 70% chars
                'zero_width': 0.6,  # 60% spaces
            }
        }

        config = rates.get(aggressiveness, rates['balanced'])

        # Strategy 1: Unicode Character Substitution
        modified = self._apply_unicode_substitution(text, config['unicode'])
        if modified != text:
            stats['changes'] += 1
            stats['techniques'].append('unicode_substitution')
            text = modified

        # Strategy 2: Zero-Width Character Insertion
        modified = self._insert_zero_width_chars(text, config['zero_width'])
        if modified != text:
            stats['changes'] += 1
            stats['techniques'].append('zero_width_insertion')
            text = modified

        # Strategy 3: Word Spacing Variation (subtle)
        modified = self._vary_word_spacing(text, config['zero_width'] * 0.5)
        if modified != text:
            stats['changes'] += 1
            stats['techniques'].append('spacing_variation')
            text = modified

        return text, stats

    def _apply_unicode_substitution(self, text: str, rate: float) -> str:
        """Replace Latin chars with visually identical Cyrillic/Greek"""
        result = []

        for char in text:
            if char in self.unicode_map and random.random() < rate:
                result.append(self.unicode_map[char])
            else:
                result.append(char)

        return ''.join(result)

    def _insert_zero_width_chars(self, text: str, rate: float) -> str:
        """Insert zero-width characters between words and within words"""
        words = text.split()
        modified_words = []

        for word in words:
            if len(word) > 2 and random.random() < rate:
                # Insert zero-width char in middle of word
                mid = len(word) // 2
                char = random.choice(self.zero_width_chars)
                word = word[:mid] + char + word[mid:]

            modified_words.append(word)

        # Also insert between words
        result = []
        for i, word in enumerate(modified_words):
            result.append(word)
            if i < len(modified_words) - 1 and random.random() < rate * 0.5:
                # Add extra zero-width char between words
                result.append(random.choice(self.zero_width_chars))

        return ' '.join(result)

    def _vary_word_spacing(self, text: str, rate: float) -> str:
        """Add subtle spacing variations using hair space"""
        # Hair space (very thin space, nearly invisible)
        hair_space = '\u200A'

        words = text.split()
        result = []

        for i, word in enumerate(words):
            result.append(word)
            if i < len(words) - 1:
                # Randomly add hair space before regular space
                if random.random() < rate:
                    result.append(hair_space)

        return ' '.join(result)

    def protect_all_headers(self, doc: docx.Document,
                          aggressiveness: str = 'balanced') -> Dict:
        """
        Protect all standard headers in document

        Args:
            doc: DOCX document
            aggressiveness: Protection level

        Returns:
            Statistics dict
        """
        stats = {
            'total_headers_found': 0,
            'headers_protected': 0,
            'total_changes': 0,
            'headers': []
        }

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if not text:
                continue

            # Check if it's a standard header
            if self.is_standard_header(text):
                stats['total_headers_found'] += 1

                # Apply protection
                protected_text, protection_stats = self.protect_header(
                    text, aggressiveness
                )

                if protection_stats['changes'] > 0:
                    # Update paragraph (preserve formatting)
                    if para.runs:
                        # Clear existing runs
                        for run in para.runs:
                            run.text = ''
                        # Set first run to protected text
                        para.runs[0].text = protected_text
                    else:
                        para.add_run(protected_text)

                    stats['headers_protected'] += 1
                    stats['total_changes'] += protection_stats['changes']

                    stats['headers'].append({
                        'paragraph_index': i,
                        'original': protection_stats['original'],
                        'protected': protected_text,
                        'techniques': protection_stats['techniques']
                    })

                    self.logger.info(f"Protected header: {text[:50]}...")

        return stats


def protect_document_headers(docx_path: str, output_path: str,
                             aggressiveness: str = 'balanced') -> Dict:
    """
    Convenience function to protect headers in a document

    Args:
        docx_path: Input DOCX path
        output_path: Output DOCX path
        aggressiveness: Protection level

    Returns:
        Protection statistics
    """
    doc = docx.Document(docx_path)
    protector = HeaderProtector()

    stats = protector.protect_all_headers(doc, aggressiveness)

    # Save document
    doc.save(output_path)

    logger = logging.getLogger(__name__)
    logger.info(f"Headers protected: {stats['headers_protected']}/{stats['total_headers_found']}")
    logger.info(f"Total changes: {stats['total_changes']}")

    return stats
