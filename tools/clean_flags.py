#!/usr/bin/env python3
"""
Clean Flags Tool
===============
Remove IPT tracking flags dari dokumen untuk submission final.

Usage:
    python tools/clean_flags.py input.docx [output.docx]
    
Options:
    --keep-properties   Keep document properties
    --dry-run          Show what would be removed without modifying

Author: DevNoLife
"""

import sys
import argparse
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def remove_hidden_markers(doc: Document) -> int:
    """Remove IPT hidden markers from document text"""
    count = 0
    
    # IPT marker patterns
    ipt_markers = [
        '\u200B\u200C\u200D',  # Main IPT marker sequence
        '\u200B[',              # Inline markers
        ']\u200C',              # End markers
    ]
    
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        cleaned_text = original_text
        
        # Remove all IPT markers
        for marker in ipt_markers:
            if marker in cleaned_text:
                cleaned_text = cleaned_text.replace(marker, '')
                count += 1
        
        # Remove marker patterns like [unicode_substitution], [zero_width_insertion]
        import re
        cleaned_text = re.sub(r'\[unicode_substitution\]', '', cleaned_text)
        cleaned_text = re.sub(r'\[zero_width_insertion\]', '', cleaned_text)
        cleaned_text = re.sub(r'\[header_modification\]', '', cleaned_text)
        cleaned_text = re.sub(r'\[metadata_change\]', '', cleaned_text)
        cleaned_text = re.sub(r'\[paraphrase\]', '', cleaned_text)
        
        if cleaned_text != original_text:
            paragraph.text = cleaned_text
    
    return count


def remove_hidden_section(doc: Document) -> bool:
    """Remove IPT MODIFICATION LOG section"""
    removed = False
    paragraphs_to_remove = []
    
    in_log_section = False
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # Check if this is the start of log section
        if 'IPT MODIFICATION LOG' in text:
            in_log_section = True
            paragraphs_to_remove.append(i)
            continue
        
        # If we're in log section, mark for removal
        if in_log_section:
            paragraphs_to_remove.append(i)
    
    # Note: python-docx doesn't provide direct paragraph removal
    # We'll clear the text instead
    for idx in paragraphs_to_remove:
        if idx < len(doc.paragraphs):
            paragraph = doc.paragraphs[idx]
            paragraph.clear()
            removed = True
    
    return removed


def remove_visual_highlights(doc: Document) -> int:
    """Remove visual highlight colors from runs"""
    count = 0
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color is not None:
                run.font.highlight_color = None
                count += 1
    
    return count


def clean_document_properties(doc: Document) -> dict:
    """Clean IPT-related document properties"""
    cleaned = {}
    
    props = doc.core_properties
    
    # Clean Subject
    if props.subject and 'IPT Modified' in props.subject:
        cleaned['subject'] = props.subject
        props.subject = ''
    
    # Clean Keywords
    if props.keywords and 'IPT' in props.keywords:
        cleaned['keywords'] = props.keywords
        props.keywords = ''
    
    # Clean Comments
    if props.comments and 'IPT' in props.comments:
        cleaned['comments'] = props.comments[:100]  # Truncate for display
        props.comments = ''
    
    return cleaned


def analyze_document(doc: Document) -> dict:
    """Analyze document for IPT artifacts"""
    analysis = {
        'has_hidden_markers': False,
        'has_visual_highlights': False,
        'has_log_section': False,
        'has_ipt_properties': False,
        'marker_count': 0,
        'highlight_count': 0
    }
    
    # Check for hidden markers
    for paragraph in doc.paragraphs:
        if '\u200B\u200C\u200D' in paragraph.text:
            analysis['has_hidden_markers'] = True
            analysis['marker_count'] += paragraph.text.count('\u200B\u200C\u200D')
        
        if 'IPT MODIFICATION LOG' in paragraph.text:
            analysis['has_log_section'] = True
    
    # Check for visual highlights
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color is not None:
                analysis['has_visual_highlights'] = True
                analysis['highlight_count'] += 1
    
    # Check properties
    props = doc.core_properties
    if props.subject and 'IPT' in props.subject:
        analysis['has_ipt_properties'] = True
    if props.keywords and 'IPT' in props.keywords:
        analysis['has_ipt_properties'] = True
    
    return analysis


def clean_document(input_path: str, output_path: str = None, 
                   keep_properties: bool = False, dry_run: bool = False) -> dict:
    """
    Clean IPT flags from document
    
    Args:
        input_path: Path to input document
        output_path: Path to output document (default: input_clean.docx)
        keep_properties: Keep document properties
        dry_run: Don't save, just report
        
    Returns:
        dict with cleaning statistics
    """
    input_file = Path(input_path)
    
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    if output_path is None:
        output_path = input_file.parent / f"{input_file.stem}_clean{input_file.suffix}"
    else:
        output_path = Path(output_path)
    
    print(f"📄 Loading document: {input_file.name}")
    doc = Document(str(input_file))
    
    # Analyze first
    print("\n🔍 Analyzing document...")
    analysis = analyze_document(doc)
    
    print(f"   Hidden markers: {'✓ Found' if analysis['has_hidden_markers'] else '✗ None'} ({analysis['marker_count']} instances)")
    print(f"   Visual highlights: {'✓ Found' if analysis['has_visual_highlights'] else '✗ None'} ({analysis['highlight_count']} runs)")
    print(f"   Log section: {'✓ Found' if analysis['has_log_section'] else '✗ None'}")
    print(f"   IPT properties: {'✓ Found' if analysis['has_ipt_properties'] else '✗ None'}")
    
    if not any([analysis['has_hidden_markers'], analysis['has_visual_highlights'], 
                analysis['has_log_section'], analysis['has_ipt_properties']]):
        print("\n✅ Document is already clean (no IPT artifacts found)")
        return {'status': 'clean', 'analysis': analysis}
    
    if dry_run:
        print("\n🏃 Dry run mode - no changes will be saved")
        return {'status': 'dry_run', 'analysis': analysis}
    
    # Perform cleaning
    print("\n🧹 Cleaning document...")
    stats = {'analysis': analysis}
    
    # 1. Remove hidden markers
    if analysis['has_hidden_markers']:
        count = remove_hidden_markers(doc)
        print(f"   ✓ Removed {count} hidden marker(s)")
        stats['markers_removed'] = count
    
    # 2. Remove visual highlights
    if analysis['has_visual_highlights']:
        count = remove_visual_highlights(doc)
        print(f"   ✓ Removed {count} visual highlight(s)")
        stats['highlights_removed'] = count
    
    # 3. Remove log section
    if analysis['has_log_section']:
        removed = remove_hidden_section(doc)
        if removed:
            print(f"   ✓ Removed change log section")
            stats['log_section_removed'] = True
    
    # 4. Clean properties
    if analysis['has_ipt_properties'] and not keep_properties:
        cleaned_props = clean_document_properties(doc)
        if cleaned_props:
            print(f"   ✓ Cleaned document properties")
            for prop, value in cleaned_props.items():
                print(f"     - {prop}: {value[:50]}...")
            stats['properties_cleaned'] = cleaned_props
    elif keep_properties:
        print(f"   ⊘ Kept document properties (--keep-properties)")
    
    # Save cleaned document
    print(f"\n💾 Saving cleaned document: {output_path.name}")
    doc.save(str(output_path))
    
    stats['status'] = 'success'
    stats['output_file'] = str(output_path)
    
    print(f"\n✅ Document cleaned successfully!")
    print(f"   Input:  {input_file}")
    print(f"   Output: {output_path}")
    
    return stats


def main():
    parser = argparse.ArgumentParser(
        description='Remove IPT tracking flags from documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python tools/clean_flags.py thesis_processed.docx
  python tools/clean_flags.py input.docx output_clean.docx
  python tools/clean_flags.py input.docx --dry-run
  python tools/clean_flags.py input.docx --keep-properties
        """
    )
    
    parser.add_argument(
        'input',
        help='Input DOCX file with IPT flags'
    )
    
    parser.add_argument(
        'output',
        nargs='?',
        default=None,
        help='Output DOCX file (default: input_clean.docx)'
    )
    
    parser.add_argument(
        '--keep-properties',
        action='store_true',
        help='Keep document properties (don\'t clean metadata)'
    )
    
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Analyze only, don\'t save changes'
    )
    
    args = parser.parse_args()
    
    try:
        result = clean_document(
            args.input,
            args.output,
            keep_properties=args.keep_properties,
            dry_run=args.dry_run
        )
        
        return 0 if result['status'] in ['success', 'clean'] else 1
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
